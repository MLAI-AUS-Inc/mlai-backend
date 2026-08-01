from io import BytesIO

from django.test import SimpleTestCase, override_settings

from docx import Document

from org_memory.drive_inventory import (
    DOCX_MIME_TYPE,
    GOOGLE_DOC_MIME_TYPE,
    PDF_MIME_TYPE,
    SRT_MIME_TYPES,
    TEXT_MIME_TYPE,
    VTT_MIME_TYPE,
)
from org_memory.drive_parsing import (
    infer_meeting_metadata,
    parse_drive_document,
)
from org_memory.models import DriveExtractionStatus, DriveWorkClassification


def text_pdf(value: str) -> bytes:
    safe = value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content = f"BT /F1 12 Tf 72 720 Td ({safe}) Tj ET".encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = []
    for number, payload in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{number} 0 obj\n".encode("ascii"))
        output.extend(payload)
        output.extend(b"\nendobj\n")
    xref = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n".encode(
            "ascii"
        )
    )
    return bytes(output)


@override_settings(
    ORG_MEMORY_DRIVE_MAX_DOWNLOAD_BYTES=1024 * 1024,
    ORG_MEMORY_DRIVE_CHUNK_TARGET_CHARS=500,
    ORG_MEMORY_DRIVE_CHUNK_MAX_CHARS=800,
    ORG_MEMORY_DRIVE_CHUNK_OVERLAP_CHARS=50,
)
class DriveParserTests(SimpleTestCase):
    def test_google_docs_markdown_preserves_headings_speakers_and_chunks(self):
        parsed = parse_drive_document(
            file_id="doc-1",
            filename="Committee Meeting 2026-07-02",
            mime_type=GOOGLE_DOC_MIME_TYPE,
            raw_bytes=(
                b"# Committee Meeting\n\n## Decisions\n\nSam: Approved the transcript pilot.\n\n"
                b"Alex: I will own the rollout by Friday."
            ),
        )

        self.assertEqual(parsed.status, DriveExtractionStatus.EXTRACTED)
        self.assertEqual(parsed.parser_name, "google_docs_markdown")
        self.assertEqual(
            [block.kind for block in parsed.blocks],
            ["heading", "heading", "speaker_turn", "speaker_turn"],
        )
        self.assertIn("Decisions", parsed.chunks[0]["source_locator"]["sections"])
        self.assertIn("Sam", parsed.chunks[0]["source_locator"]["speakers"])

    def test_docx_preserves_heading_paragraph_and_table_locations(self):
        document = Document()
        document.add_heading("Board Meeting", level=1)
        document.add_paragraph("Sam: Approved the budget.")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Owner"
        table.cell(0, 1).text = "Alex"
        output = BytesIO()
        document.save(output)

        parsed = parse_drive_document(
            file_id="docx-1",
            filename="Board Meeting.docx",
            mime_type=DOCX_MIME_TYPE,
            raw_bytes=output.getvalue(),
        )

        self.assertEqual(parsed.status, DriveExtractionStatus.EXTRACTED)
        self.assertEqual(parsed.blocks[0].kind, "heading")
        self.assertEqual(parsed.blocks[1].speaker, "Sam")
        self.assertEqual(parsed.blocks[-1].kind, "table_row")
        self.assertEqual(parsed.blocks[-1].locator["table"], 1)

    def test_pdf_has_page_locator_and_blank_pdf_is_visible_ocr_work(self):
        parsed = parse_drive_document(
            file_id="pdf-1",
            filename="Town Hall Transcript.pdf",
            mime_type=PDF_MIME_TYPE,
            raw_bytes=text_pdf("Sam: Approved the pilot."),
        )
        self.assertEqual(parsed.status, DriveExtractionStatus.EXTRACTED)
        self.assertEqual(parsed.blocks[0].locator["page"], 1)

        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=612, height=792)
        blank = BytesIO()
        writer.write(blank)
        scanned = parse_drive_document(
            file_id="pdf-2",
            filename="Scanned Meeting.pdf",
            mime_type=PDF_MIME_TYPE,
            raw_bytes=blank.getvalue(),
        )
        self.assertEqual(scanned.status, DriveExtractionStatus.UNSUPPORTED)
        self.assertEqual(scanned.work_classification, DriveWorkClassification.NEEDS_OCR)
        self.assertIn("no_extractable_text", scanned.warnings)

    def test_webvtt_preserves_speaker_and_timestamp_ranges(self):
        parsed = parse_drive_document(
            file_id="vtt-1",
            filename="Standup 2026-07-03.vtt",
            mime_type=VTT_MIME_TYPE,
            raw_bytes=(
                b"WEBVTT\n\n00:00:01.000 --> 00:00:04.000\nSam: The pilot is approved.\n\n"
                b"00:00:05.000 --> 00:00:08.000\nAlex: I will ship it Friday.\n"
            ),
        )
        self.assertEqual(parsed.status, DriveExtractionStatus.EXTRACTED)
        self.assertEqual(parsed.blocks[0].speaker, "Sam")
        self.assertEqual(parsed.blocks[0].locator["timestamp_start_seconds"], 1.0)
        self.assertEqual(parsed.chunks[0]["source_locator"]["timestamp_end"], "00:00:04.000")

        srt = parse_drive_document(
            file_id="srt-1",
            filename="Standup.srt",
            mime_type=next(iter(SRT_MIME_TYPES)),
            raw_bytes=b"1\n00:00:01,000 --> 00:00:03,000\nSam: SRT is preserved.\n",
        )
        self.assertEqual(srt.status, DriveExtractionStatus.EXTRACTED)
        self.assertEqual(srt.blocks[0].speaker, "Sam")
        self.assertEqual(srt.blocks[0].locator["cue"], 1)

        plain = parse_drive_document(
            file_id="txt-1",
            filename="Meeting transcript.txt",
            mime_type=TEXT_MIME_TYPE,
            raw_bytes=b"Sam: Plain text is preserved.",
        )
        self.assertEqual(plain.status, DriveExtractionStatus.EXTRACTED)
        self.assertEqual(plain.parser_name, "plain_text")

    def test_meeting_identity_is_stable_for_copied_filename(self):
        original = infer_meeting_metadata(
            filename="Committee Meeting Transcript 2026-07-02.md",
            text="Sam: Approved the pilot.",
        )
        copied = infer_meeting_metadata(
            filename="Committee Meeting 2026-07-02 - Copy.md",
            text="Sam: Approved the pilot.",
        )
        self.assertEqual(original["identity_key"], copied["identity_key"])
        self.assertEqual(original["occurred_at"][:10], "2026-07-02")
        self.assertEqual(original["participants"], ["Sam"])

    def test_meeting_datetime_preserves_time_from_title(self):
        metadata = infer_meeting_metadata(
            filename="MLAI Committee Meeting – 2026/07/20 18:30 AEST – Notes by Gemini",
            text="The committee approved the plan.",
        )

        self.assertEqual(metadata["occurred_at"], "2026-07-20T18:30:00+10:00")
        self.assertEqual(
            metadata["identity_basis"]["date_basis"],
            "filename_or_heading_datetime",
        )
        daylight = infer_meeting_metadata(
            filename="MLAI Committee Meeting – 2026/12/14 18:30 AEDT – Notes by Gemini",
            text="The committee approved the summer plan.",
        )
        self.assertEqual(daylight["occurred_at"], "2026-12-14T18:30:00+11:00")
