from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass, field
from datetime import datetime, time as datetime_time
from pathlib import PurePath
from typing import Iterable, Mapping, Optional
from zoneinfo import ZoneInfo

from django.conf import settings

from .drive_inventory import (
    DOCX_MIME_TYPE,
    GOOGLE_DOC_MIME_TYPE,
    MARKDOWN_MIME_TYPE,
    PDF_MIME_TYPE,
    SRT_MIME_TYPES,
    TEXT_MIME_TYPE,
    VTT_MIME_TYPE,
)
from .models import DriveExtractionStatus, DriveWorkClassification


DRIVE_PARSER_VERSION = "drive-parser-v1"
GOOGLE_DOC_EXPORT_MIME_TYPE = MARKDOWN_MIME_TYPE

TIMESTAMP_PATTERN = re.compile(
    r"(?P<start>\d{1,2}:\d{2}(?::\d{2})?(?:[.,]\d{1,3})?)\s*--?>\s*"
    r"(?P<end>\d{1,2}:\d{2}(?::\d{2})?(?:[.,]\d{1,3})?)"
)
SPEAKER_PATTERN = re.compile(
    r"^(?:\[[^\]]{1,32}\]\s*)?(?P<speaker>[A-Za-z][A-Za-z0-9 ._'’-]{0,79}):\s+(?P<text>.+)$"
)
DATE_PATTERNS = (
    re.compile(r"(?<!\d)(?P<year>20\d{2})[-_/](?P<month>0?[1-9]|1[0-2])[-_/](?P<day>0?[1-9]|[12]\d|3[01])(?!\d)"),
    re.compile(
        r"(?<!\d)(?P<day>0?[1-9]|[12]\d|3[01])\s+"
        r"(?P<month_name>Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|"
        r"Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|"
        r"Nov(?:ember)?|Dec(?:ember)?)\s+(?P<year>20\d{2})(?!\d)",
        re.IGNORECASE,
    ),
)
MONTHS = {
    "jan": 1,
    "feb": 2,
    "mar": 3,
    "apr": 4,
    "may": 5,
    "jun": 6,
    "jul": 7,
    "aug": 8,
    "sep": 9,
    "oct": 10,
    "nov": 11,
    "dec": 12,
}
TITLE_NOISE_PATTERN = re.compile(
    r"\b(transcript|meeting notes?|minutes|recording|copy)\b|\(\d+\)$",
    re.IGNORECASE,
)
SPACE_PATTERN = re.compile(r"\s+")


class DriveParseError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedBlock:
    text: str
    kind: str = "paragraph"
    locator: Mapping = field(default_factory=dict)
    speaker: str = ""
    timestamp_start: str = ""
    timestamp_end: str = ""


@dataclass(frozen=True)
class ParsedDriveDocument:
    status: str
    parser_name: str
    text: str
    blocks: tuple[ParsedBlock, ...]
    chunks: tuple[Mapping, ...]
    warnings: tuple[str, ...]
    work_classification: str = DriveWorkClassification.NONE
    error: str = ""


def _clean_text(value: str) -> str:
    value = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    value = "\n".join(line.rstrip() for line in value.split("\n"))
    return re.sub(r"\n{3,}", "\n\n", value).strip()


def _speaker_parts(text: str) -> tuple[str, str]:
    match = SPEAKER_PATTERN.match(text.strip())
    if not match:
        return "", text.strip()
    return match.group("speaker").strip(), match.group("text").strip()


def _parse_markdown_or_text(raw_bytes: bytes, *, markdown: bool) -> tuple[list[ParsedBlock], list[str]]:
    decoded = raw_bytes.decode("utf-8", errors="replace")
    warnings = ["invalid_utf8_replaced"] if "\ufffd" in decoded else []
    lines = decoded.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks = []
    paragraph = []
    paragraph_start = 1
    section = ""

    def flush(end_line: int):
        nonlocal paragraph
        text = _clean_text("\n".join(paragraph))
        if text:
            speaker, body = _speaker_parts(text)
            blocks.append(
                ParsedBlock(
                    text=f"{speaker}: {body}" if speaker else body,
                    kind="speaker_turn" if speaker else "paragraph",
                    speaker=speaker,
                    locator={
                        "line_start": paragraph_start,
                        "line_end": end_line,
                        "section": section,
                    },
                )
            )
        paragraph = []

    for line_number, line in enumerate(lines, start=1):
        stripped = line.strip()
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped) if markdown else None
        if heading:
            flush(line_number - 1)
            section = heading.group(2).strip()
            blocks.append(
                ParsedBlock(
                    text=section,
                    kind="heading",
                    locator={"line_start": line_number, "line_end": line_number, "section": section},
                )
            )
        elif SPEAKER_PATTERN.match(stripped):
            flush(line_number - 1)
            speaker, body = _speaker_parts(stripped)
            blocks.append(
                ParsedBlock(
                    text=f"{speaker}: {body}",
                    kind="speaker_turn",
                    speaker=speaker,
                    locator={
                        "line_start": line_number,
                        "line_end": line_number,
                        "section": section,
                    },
                )
            )
            paragraph_start = line_number + 1
        elif not stripped:
            flush(line_number - 1)
            paragraph_start = line_number + 1
        else:
            if not paragraph:
                paragraph_start = line_number
            paragraph.append(stripped)
    flush(len(lines))
    return blocks, warnings


def _parse_docx(raw_bytes: bytes) -> tuple[list[ParsedBlock], list[str]]:
    from docx import Document

    document = Document(io.BytesIO(raw_bytes))
    blocks = []
    section = ""
    for index, paragraph in enumerate(document.paragraphs):
        text = _clean_text(paragraph.text)
        if not text:
            continue
        style = str(getattr(paragraph.style, "name", "") or "")
        is_heading = style.lower().startswith("heading")
        if is_heading:
            section = text
        speaker, body = _speaker_parts(text)
        blocks.append(
            ParsedBlock(
                text=f"{speaker}: {body}" if speaker else body,
                kind="heading" if is_heading else "speaker_turn" if speaker else "paragraph",
                speaker=speaker,
                locator={"paragraph": index + 1, "section": section, "style": style[:80]},
            )
        )
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            text = " | ".join(_clean_text(cell.text) for cell in row.cells if _clean_text(cell.text))
            if text:
                blocks.append(
                    ParsedBlock(
                        text=text,
                        kind="table_row",
                        locator={"table": table_index, "row": row_index, "section": section},
                    )
                )
    return blocks, []


def _parse_pdf(raw_bytes: bytes) -> tuple[list[ParsedBlock], list[str], str]:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(raw_bytes))
    blocks = []
    warnings = []
    empty_pages = 0
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = _clean_text(page.extract_text() or "")
        if not page_text:
            empty_pages += 1
            continue
        for paragraph_index, paragraph in enumerate(re.split(r"\n\s*\n|\n", page_text), start=1):
            text = _clean_text(paragraph)
            if not text:
                continue
            speaker, body = _speaker_parts(text)
            blocks.append(
                ParsedBlock(
                    text=f"{speaker}: {body}" if speaker else body,
                    kind="speaker_turn" if speaker else "paragraph",
                    speaker=speaker,
                    locator={"page": page_number, "paragraph": paragraph_index},
                )
            )
    if empty_pages:
        warnings.append(f"pdf_pages_without_text:{empty_pages}")
    work = DriveWorkClassification.NEEDS_OCR if not blocks else DriveWorkClassification.NONE
    return blocks, warnings, work


def _seconds(timestamp: str) -> float:
    parts = timestamp.replace(",", ".").split(":")
    if len(parts) == 2:
        hours = 0
        minutes, seconds = parts
    else:
        hours, minutes, seconds = parts
    return int(hours) * 3600 + int(minutes) * 60 + float(seconds)


def _parse_cues(raw_bytes: bytes, *, vtt: bool) -> tuple[list[ParsedBlock], list[str]]:
    text = raw_bytes.decode("utf-8-sig", errors="replace").replace("\r\n", "\n").replace("\r", "\n")
    warnings = ["invalid_utf8_replaced"] if "\ufffd" in text else []
    groups = re.split(r"\n\s*\n", text.strip())
    blocks = []
    cue_number = 0
    for group in groups:
        lines = [line.strip() for line in group.split("\n") if line.strip()]
        if not lines or (vtt and lines[0].upper().startswith("WEBVTT")):
            continue
        timestamp_index = next(
            (index for index, line in enumerate(lines) if TIMESTAMP_PATTERN.search(line)),
            None,
        )
        if timestamp_index is None:
            warnings.append("caption_block_without_timestamp")
            continue
        match = TIMESTAMP_PATTERN.search(lines[timestamp_index])
        cue_text = _clean_text(" ".join(lines[timestamp_index + 1 :]))
        cue_text = re.sub(r"<[^>]+>", "", cue_text).strip()
        if not cue_text:
            continue
        cue_number += 1
        speaker, body = _speaker_parts(cue_text)
        blocks.append(
            ParsedBlock(
                text=f"{speaker}: {body}" if speaker else body,
                kind="speaker_turn" if speaker else "caption_cue",
                speaker=speaker,
                timestamp_start=match.group("start"),
                timestamp_end=match.group("end"),
                locator={
                    "cue": cue_number,
                    "timestamp_start_seconds": _seconds(match.group("start")),
                    "timestamp_end_seconds": _seconds(match.group("end")),
                },
            )
        )
    return blocks, warnings


def _split_long_block(block: ParsedBlock, max_chars: int, overlap: int) -> Iterable[ParsedBlock]:
    if len(block.text) <= max_chars:
        yield block
        return
    start = 0
    part = 1
    while start < len(block.text):
        end = min(start + max_chars, len(block.text))
        if end < len(block.text):
            boundary = block.text.rfind(" ", start + max_chars // 2, end)
            if boundary > start:
                end = boundary
        locator = dict(block.locator)
        locator.update({"part": part, "part_start": start, "part_end": end})
        yield ParsedBlock(
            text=block.text[start:end].strip(),
            kind=block.kind,
            locator=locator,
            speaker=block.speaker,
            timestamp_start=block.timestamp_start,
            timestamp_end=block.timestamp_end,
        )
        if end >= len(block.text):
            break
        start = max(end - overlap, start + 1)
        part += 1


def build_chunks(blocks: Iterable[ParsedBlock], *, file_id: str, parser_name: str) -> tuple[dict, ...]:
    target = max(int(settings.ORG_MEMORY_DRIVE_CHUNK_TARGET_CHARS), 500)
    maximum = max(int(settings.ORG_MEMORY_DRIVE_CHUNK_MAX_CHARS), target)
    overlap = min(max(int(settings.ORG_MEMORY_DRIVE_CHUNK_OVERLAP_CHARS), 0), maximum // 4)
    expanded = [part for block in blocks for part in _split_long_block(block, maximum, overlap)]
    chunks = []
    pending = []
    pending_size = 0
    pending_start = 0
    pending_end = 0

    def flush():
        nonlocal pending, pending_size, pending_start, pending_end
        if not pending:
            return
        text = "\n\n".join(block.text for block in pending).strip()
        speakers = sorted({block.speaker for block in pending if block.speaker})
        sections = []
        pages = []
        for block in pending:
            section = block.locator.get("section")
            page = block.locator.get("page")
            if section and section not in sections:
                sections.append(section)
            if page is not None:
                pages.append(page)
        chunks.append(
            {
                "ordinal": len(chunks),
                "text": text,
                "token_count": max(round(len(text) / 4), 1),
                "chunk_kind": (
                    "speaker_turn"
                    if all(block.kind in {"speaker_turn", "caption_cue"} for block in pending)
                    else "document_section"
                ),
                "source_locator": {
                    "file_id": file_id,
                    "parser": parser_name,
                    "block_start": pending_start,
                    "block_end": pending_end,
                    "sections": sections[:20],
                    "page_start": min(pages) if pages else None,
                    "page_end": max(pages) if pages else None,
                    "speakers": speakers[:50],
                    "timestamp_start": pending[0].timestamp_start or None,
                    "timestamp_end": pending[-1].timestamp_end or None,
                },
                "start_offset": None,
                "end_offset": None,
            }
        )
        pending = []
        pending_size = 0

    for block_index, block in enumerate(expanded):
        block_size = len(block.text) + (2 if pending else 0)
        timestamped = bool(block.timestamp_start or block.timestamp_end)
        if pending and (pending_size + block_size > target or timestamped):
            flush()
        if not pending:
            pending_start = block_index
        pending.append(block)
        pending_end = block_index
        pending_size += block_size
        if timestamped or pending_size >= target:
            flush()
    flush()
    return tuple(chunks)


def parse_drive_document(
    *,
    file_id: str,
    filename: str,
    mime_type: str,
    raw_bytes: bytes,
) -> ParsedDriveDocument:
    maximum = max(int(settings.ORG_MEMORY_DRIVE_MAX_DOWNLOAD_BYTES), 1)
    if len(raw_bytes) > maximum:
        return ParsedDriveDocument(
            status=DriveExtractionStatus.UNSUPPORTED,
            parser_name="size_guard",
            text="",
            blocks=(),
            chunks=(),
            warnings=(f"download_exceeds_limit:{maximum}",),
            work_classification=DriveWorkClassification.UNSUPPORTED_FORMAT,
        )
    extension = PurePath(filename or "").suffix.lower()
    mime_type = str(mime_type or "").lower()
    try:
        if mime_type == GOOGLE_DOC_MIME_TYPE:
            blocks, warnings = _parse_markdown_or_text(raw_bytes, markdown=True)
            parser_name = "google_docs_markdown"
            work = DriveWorkClassification.NONE
        elif mime_type == DOCX_MIME_TYPE or extension == ".docx":
            blocks, warnings = _parse_docx(raw_bytes)
            parser_name = "docx"
            work = DriveWorkClassification.NONE
        elif mime_type == PDF_MIME_TYPE or extension == ".pdf":
            blocks, warnings, work = _parse_pdf(raw_bytes)
            parser_name = "pdf_text_layer"
        elif mime_type == VTT_MIME_TYPE or extension == ".vtt":
            blocks, warnings = _parse_cues(raw_bytes, vtt=True)
            parser_name = "webvtt"
            work = DriveWorkClassification.NONE
        elif mime_type in SRT_MIME_TYPES or extension == ".srt":
            blocks, warnings = _parse_cues(raw_bytes, vtt=False)
            parser_name = "srt"
            work = DriveWorkClassification.NONE
        elif mime_type in {TEXT_MIME_TYPE, MARKDOWN_MIME_TYPE} or extension in {
            ".txt",
            ".md",
        }:
            is_markdown = extension == ".md" or mime_type == MARKDOWN_MIME_TYPE
            blocks, warnings = _parse_markdown_or_text(
                raw_bytes,
                markdown=is_markdown,
            )
            parser_name = "markdown" if is_markdown else "plain_text"
            work = DriveWorkClassification.NONE
        else:
            return ParsedDriveDocument(
                status=DriveExtractionStatus.UNSUPPORTED,
                parser_name="unsupported",
                text="",
                blocks=(),
                chunks=(),
                warnings=(f"unsupported_mime_type:{mime_type or 'unknown'}",),
                work_classification=DriveWorkClassification.UNSUPPORTED_FORMAT,
            )
    except Exception as exc:
        return ParsedDriveDocument(
            status=DriveExtractionStatus.FAILED,
            parser_name="parse_error",
            text="",
            blocks=(),
            chunks=(),
            warnings=(f"parse_error:{type(exc).__name__}",),
            error=str(exc)[:1000],
        )
    text = _clean_text("\n\n".join(block.text for block in blocks))
    if not text:
        classification = work or DriveWorkClassification.UNSUPPORTED_FORMAT
        return ParsedDriveDocument(
            status=DriveExtractionStatus.UNSUPPORTED,
            parser_name=parser_name,
            text="",
            blocks=tuple(blocks),
            chunks=(),
            warnings=tuple(sorted(set(warnings) | {"no_extractable_text"})),
            work_classification=classification,
        )
    chunks = build_chunks(blocks, file_id=file_id, parser_name=parser_name)
    return ParsedDriveDocument(
        status=DriveExtractionStatus.EXTRACTED,
        parser_name=parser_name,
        text=text,
        blocks=tuple(blocks),
        chunks=chunks,
        warnings=tuple(sorted(set(warnings))),
    )


def normalized_content(value: str) -> str:
    return SPACE_PATTERN.sub(" ", str(value or "").lower()).strip()


def content_signature(value: str, *, size: int = 64) -> list[str]:
    words = normalized_content(value).split()
    shingles = {
        hashlib.sha256(" ".join(words[index : index + 5]).encode("utf-8")).hexdigest()[:16]
        for index in range(max(len(words) - 4, 1))
        if words[index : index + 5]
    }
    return sorted(shingles)[:size]


def normalized_title(filename: str) -> str:
    title = PurePath(str(filename or "")).stem
    for pattern in DATE_PATTERNS:
        title = pattern.sub(" ", title)
    title = TITLE_NOISE_PATTERN.sub(" ", title)
    return SPACE_PATTERN.sub(" ", re.sub(r"[^A-Za-z0-9]+", " ", title)).strip().lower() or "untitled meeting"


def infer_meeting_metadata(
    *,
    filename: str,
    text: str,
    source_created_at: Optional[str] = None,
    timezone_name: str = "Australia/Sydney",
) -> dict:
    title = normalized_title(filename)
    search_value = f"{filename}\n{text[:4000]}"
    occurred_at = None
    date_basis = "unknown"
    for pattern in DATE_PATTERNS:
        match = pattern.search(search_value)
        if not match:
            continue
        values = match.groupdict()
        month = values.get("month") or MONTHS[values["month_name"][:3].lower()]
        try:
            occurred_at = datetime.combine(
                datetime(int(values["year"]), int(month), int(values["day"])).date(),
                datetime_time.min,
                tzinfo=ZoneInfo(timezone_name),
            )
        except ValueError:
            occurred_at = None
        if occurred_at:
            date_basis = "filename_or_heading"
            break
    speakers = []
    for line in text.splitlines():
        speaker, _body = _speaker_parts(line)
        if speaker and speaker not in speakers:
            speakers.append(speaker)
    identity_payload = {
        "title": title,
        "date": occurred_at.date().isoformat() if occurred_at else None,
    }
    if not occurred_at:
        identity_payload["source_created_date"] = str(source_created_at or "")[:10]
    identity_key = hashlib.sha256(
        repr(sorted(identity_payload.items())).encode("utf-8")
    ).hexdigest()
    return {
        "identity_key": identity_key,
        "normalized_title": title,
        "occurred_at": occurred_at.isoformat() if occurred_at else None,
        "timezone_name": timezone_name,
        "participants": speakers[:100],
        "identity_basis": {**identity_payload, "date_basis": date_basis},
    }
